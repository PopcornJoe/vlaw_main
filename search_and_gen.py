import streamlit as st
import pandas as pd
from docx import Document
import datetime
import os
import pyodbc

def get_connection():
    connection_string = (
        "DRIVER={ODBC Driver 17 for SQL Server};"
        "SERVER=" + st.secrets["database"]["server"] + ";"
        "DATABASE=" + st.secrets["database"]["database"] + ";"
        "UID=" + st.secrets["database"]["user"] + ";"
        "PWD=" + st.secrets["database"]["password"] + ";"
        "Connection Timeout=60;"
        
    )
    return pyodbc.connect(connection_string)

def flatten_row(row, expected_length):
    """
    Convert a fetched row into a plain list of the expected length.
    
    - First, force the row to a list.
    - If the resulting list has a single element that is itself a list or tuple with the expected length,
      then unwrap that inner element.
    - Otherwise, return the list.
    """
    new_row = list(row)  # Force conversion to a list
    if len(new_row) == 1:
        inner = new_row[0]
        if isinstance(inner, (list, tuple)) and len(inner) == expected_length:
            return list(inner)
    return new_row

def search_matter(matter_no):
    conn = get_connection()
    cur = conn.cursor()
    query = "SELECT * FROM vlaw_base WHERE mat = ?"
    cur.execute(query, (matter_no,))
    rows = cur.fetchall()
    colnames = [desc[0] for desc in cur.description]
    cur.close()
    conn.close()
    
    # Process each row using flatten_row.
    flattened_rows = [flatten_row(row, len(colnames)) for row in rows]
    
    # Debug prints (optional)
    if flattened_rows:
        print("After flattening, first row:", flattened_rows[0])
        print("Length of first row:", len(flattened_rows[0]))
    
    return pd.DataFrame(flattened_rows, columns=colnames)





# ----------------------
# Template Functions
# ----------------------
def get_template_choices():
    """
    Returns a dictionary mapping template names to file paths.
    Make sure these file paths match your actual folder structure.
    """
    return {
        "Cost Letter": "data/templates/cost_letter.docx",
        "Demand FNB to FMC": "data/templates/demand_fnb_to_fmc.docx",
        "LOD FMC to Debtor": "data/templates/lod_fmc_to_debtor.docx",
        "Regulation 19": "data/templates/regulation_19.docx",
        "Section 129": "data/templates/section_129.docx",
        "Section 129 and Regulation 19": "data/templates/section_129_and_regulation_19.docx",
        "Defendant Summons": "data/templates/1 Defendant Summons 2025.docx"
    }

# Define which fields (template tags) are needed for each template.
template_fields = {
    "Cost Letter": [
        "debtor_1", "address", "email", "current_date",
        "matter_number", "MATTER_NAME", "account_number", "arrears_amount", "debtor_2"
    ],
    "Demand FNB to FMC": [
        "current_date", "debtor_1", "debtor_2", "account_number", "balance_amount"
    ],
    "LOD FMC to Debtor": [
        "debtor_1", "address", "email", "current_date", "matter_number",
        "instalment_amount", "arrears_amount", "balance_amount", "property_description"
    ],
    "Regulation 19": [
        "debtor_1", "address", "email", "current_date", "title",
        "matter_name", "account_number", "instalment_amount",
        "arrears_amount", "balance_amount", "property_description"
    ],
    "Section 129": [
        "debtor_1", "address", "email", "current_date", "title",
        "matter_name", "account_number", "instalment_amount",
        "arrears_amount", "balance_amount", "property_description"
    ],
    "Section 129 and Regulation 19": [
        "debtor_1", "address", "email", "current_date", "title",
        "matter_name", "account_number", "instalment_amount",
        "arrears_amount", "balance_amount", "property_description"
    ],
    "Defendant Summons": [
        # Header section tags
        "fnbaccno", "MAT", "Division", "Court",
        "Defname1",  "ID1", "Domadd", "immadd", "sumdate", "caseno",
        # Loan and bond details
        "Loandate1", "loanplace1", "loandate2", "loanplace2",
        "bondno1", "bondregdate1", "bondno2", "bondregdate2",
        "loanamount1", "loanamount2",
        # Payment and interest details
        "LOANINTRATE", "INSTALMENTPERIOD1", "INSTALMENTAMOUNT1",
        "INSTALMENTPERIOD2", "INSTALMENTOUT2", "INSTALMENTAMOUNT",
        # Certificate and arrears details
        "CERTDATE", "RCAPITAL", "JINTRATE", "JDATE",
        "ARREARSAMOUNT", "STATEMENTDATES",
        # Additional bond amounts
        "BONDAMOUNT", "ADDAMOUNT", "bondamount2", "addamount2"
    ]
}

# Mapping of database column names to the template tag names.
db_to_template_mapping = {
    "mat": "MAT",
    "defname1": "debtor_1",
    "immadd": "address",
    "email": "email",
    "mat":"matter_number",
    "matter_name": "matter_name",
    "fnbaccno": "account_number",
    "arrearsamount": "arrears_amount",
    "instalmentamount": "instalment_amount",
    "rcapital": "balance_amount",
    "defname2":"debtor_2",
    "title": "title",
    "fnbaccno":"fnbaccno",
    "defname1":"Defname1"
    

}

def render_template(template_path, data, output_path):
    """
    Loads the docx template, replaces placeholders with values from data,
    and saves the modified document.
    
    Placeholders in the document should be in the format <<key>>.
    """
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"<<{key}>>"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        
    doc.save(output_path)

# ----------------------
# Main Application
# ----------------------
def app():
    st.title("Document Generator for Matters")

    # --- Matter Search Section ---
    st.header("Search for Matter")
    matter_no = st.text_input("Enter Matter Number:")
    if st.button("Search Matter"):
        if matter_no:
            df = search_matter(matter_no)
            if df.empty:
                st.error(f"No record found for matter number: {matter_no}")
            else:
                st.success("Record found!")
                st.dataframe(df)
                # Store the first matching record in session state.
                st.session_state["record"] = df.iloc[0].to_dict()
        else:
            st.error("Please enter a matter number.")

    # --- Template Selection and Form Section ---
    if "record" in st.session_state:
        st.header("Select Template and Auto-Populate Form")
        template_choices = get_template_choices()
        selected_template = st.selectbox("Choose a Template", list(template_choices.keys()))
        template_file = template_choices[selected_template]
        
        # Get the required template fields for the selected template.
        req_fields = template_fields.get(selected_template, [])
        
        st.subheader("Template Data (Edit if needed)")
        context = {}
        
        # For each required field, attempt to auto-populate from the database record.
        for field in req_fields:
            # If the field is "current_date", pre-fill with today's date.
            if field == "current_date":
                default_val = datetime.date.today().strftime("%Y-%m-%d")
            else:
                # Check if there's a matching value from the record.
                default_val = ""
                for db_field, tmpl_field in db_to_template_mapping.items():
                    if tmpl_field == field and db_field in st.session_state["record"]:
                        default_val = st.session_state["record"][db_field]
                        break
            context[field] = st.text_input(field, value=str(default_val))
        
        # --- Generate and Download Document ---
        if st.button("Generate Document"):
            # Simple validation: ensure all required fields are filled.
            if all(context.values()):
                output_filename = f"generated_{selected_template.replace(' ', '_')}.docx"
                render_template(template_file, context, output_filename)
                st.success(f"Document generated: {output_filename}")
                
                # Provide a download button
                with open(output_filename, "rb") as f:
                    st.download_button("Download Document", data=f, file_name=output_filename)
            else:
                st.error("Please complete all fields.")
