# src/word_processor.py
import os
from docx import Document

TEMPLATES = {
    'section_129': os.path.join("data", "templates", "section_129.docx"),
    'regulation_19': os.path.join("data", "templates", "regulation_19.docx"),
    'cost_letter': os.path.join("data", "templates", "cost_letter.docx"),
    'section_129_and_regulation_19': os.path.join("data", "templates", "section_129_and_regulation_19.docx"),
    'demand_fnb_to_fmc': os.path.join("data", "templates", "demand_fnb_to_fmc.docx"),
    'lod_fmc_to_debtor': os.path.join("data", "templates", "lod_fmc_to_debtor.docx")
}

def populate_word_document(template_path, data):
    '''
    Populates a Word document with data.
    '''
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"<<{key}>>" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"<<{key}>>", str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"<<{key}>>" in cell.text:
                        cell.text = cell.text.replace(f"<<{key}>>", str(value))
    
    return doc

def process_templates(selected_templates, data):
    '''
    Loops through the selected templates and generates populated Word documents.
    '''
    generated_docs = {}
    for template_name in selected_templates:
        if template_name in TEMPLATES:
            doc = populate_word_document(TEMPLATES[template_name], data)
            output_path = os.path.join("output", f"{template_name}_{data['matter_number']}.docx")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            generated_docs[template_name] = output_path
    return generated_docs